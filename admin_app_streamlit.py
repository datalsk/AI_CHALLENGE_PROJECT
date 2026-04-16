import streamlit as st
import pandas as pd
import boto3
import json
import urllib.parse
from datetime import datetime
import requests
import io
import os

import openpyxl
from openpyxl.styles import Alignment, Font, Border, Side, PatternFill, Protection
from openpyxl.drawing.image import Image as ExcelImage
from PIL import Image
import docx
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ==========================================
# 0. UI 설정 및 컴팩트 SaaS 디자인 CSS 적용
# ==========================================
st.set_page_config(page_title="관리자 대시보드", layout="wide")

st.markdown("""
    <style>
    @import url("https://cdn.jsdelivr.net/gh/orioncactus/pretendard@v1.3.8/dist/web/variable/pretendardvariable.css");
    html, body, [class*="css"], .stMarkdown, .stText, button, input, select, .stDataFrame {
        font-family: 'Pretendard Variable', Pretendard, -apple-system, sans-serif !important;
    }
    
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    .stAppDeployButton {display:none;}
    header {background-color: transparent !important;}
    
    /* 우측 상단 툴바 완벽 숨김 */
    [data-testid="stHeaderActionElements"] {display: none !important;}
    [data-testid="stToolbar"] {visibility: hidden !important;}

    /* 카드 패딩 극한으로 축소 */
    div[data-testid="stVerticalBlockBorderWrapper"] {
        border-radius: 8px !important;
        box-shadow: rgba(0, 0, 0, 0.02) 0px 2px 4px !important;
        border: 1px solid rgba(148, 163, 184, 0.2) !important;
        background-color: rgba(255, 255, 255, 0.02) !important;
        padding: 4px 8px !important; 
        margin-bottom: 0px !important;
        transition: all 0.2s ease;
    }
    
    /* 기본 컬럼 갭 축소 */
    [data-testid="column"] > div {
        gap: 0.3rem !important;
    }
    
    /* 주요 버튼 */
    .stButton > button[kind="primary"] {
        background-color: #4f46e5;
        color: white;
        border: none;
        border-radius: 8px;
        font-weight: 600;
        letter-spacing: -0.3px;
        transition: all 0.2s ease;
    }
    .stButton > button[kind="primary"]:hover {
        background-color: #3730a3;
        box-shadow: 0 4px 12px rgba(79, 70, 229, 0.3);
    }
    
    /* 이모지 버튼 정중앙 정렬 및 여백 최소화 */
    .stButton > button[kind="secondary"], div[data-testid="stPopover"] > button {
        border-radius: 6px !important;
        font-weight: 600 !important;
        font-size: 14px !important; 
        border: 1px solid rgba(148, 163, 184, 0.3) !important;
        background-color: transparent !important;
        padding: 0px 12px !important;
        height: 36px !important;
        min-height: 36px !important; 
        display: flex !important;
        align-items: center !important;
        justify-content: center !important;
        margin: 0 !important;
        white-space: nowrap !important;
    }
    .stButton > button[kind="secondary"]:hover, div[data-testid="stPopover"] > button:hover {
        background-color: rgba(148, 163, 184, 0.1) !important;
    }
    
    h1 { font-weight: 700 !important; letter-spacing: -1px; margin-bottom: 0px !important;}
    h3 { font-weight: 600 !important; letter-spacing: -0.5px; }
    
    /* 입력 폼 텍스트 잘림 방지 및 높이 통일 */
    div[data-baseweb="input"], div[data-baseweb="select"] {
        border-radius: 6px !important;
        border: none !important;
        background-color: rgba(148, 163, 184, 0.08) !important;
        box-shadow: inset 0 0 0 1px rgba(148, 163, 184, 0.2) !important;
        transition: all 0.2s ease;
        height: 36px !important;
        min-height: 36px !important; 
    }
    div[data-baseweb="input"]:focus-within, div[data-baseweb="select"]:focus-within {
        box-shadow: inset 0 0 0 2px #4f46e5 !important;
        background-color: rgba(148, 163, 184, 0.12) !important;
    }
    
    /* 일반 텍스트 입력창 여백 */
    div[data-baseweb="input"] > div > input {
        background-color: transparent !important;
        padding-left: 8px !important; 
        padding-right: 8px !important;
        padding-top: 2px !important;
        padding-bottom: 2px !important;
        font-size: 14px !important;
    }
    
    /* 드롭다운(Select) 내부 패딩을 0에 가깝게 줄여 글자 확보 */
    div[data-baseweb="select"] > div {
        background-color: transparent !important;
        padding-left: 4px !important; 
        padding-right: 0px !important;
        padding-top: 2px !important;
        padding-bottom: 2px !important;
        font-size: 14px !important;
    }
    </style>
    """, unsafe_allow_html=True)

# ==========================================
# 1. 설정, S3 연동 및 [팀별 고정 명단]
# ==========================================
try:
    AWS_ACCESS_KEY_ID = st.secrets["AWS_ACCESS_KEY_ID"]
    AWS_SECRET_ACCESS_KEY = st.secrets["AWS_SECRET_ACCESS_KEY"]
    S3_BUCKET_NAME = st.secrets["S3_BUCKET_NAME"]
    AWS_REGION = st.secrets["AWS_REGION"]
except:
    st.error("설정 파일(secrets.toml)이 없습니다. 로컬 테스트 시 설정을 확인하세요.")
    st.stop()

FIXED_CATEGORIES = ["야근식대", "야근교통비", "외근교통비", "프로젝트비용", "기타"]

TEAM_MEMBERS = {
    "DX2본부": [
        "송은주", "이재상", "이한새", "강윤희", "김정란", "이현지", 
        "김광수", "정창호", "박형규", "김정민", "김지현", "최보영", 
        "이은진", "손민우", "오상윤", "구대연", "이승규", "임성묵", 
        "최우혁", "박주연", "안소희", "구재현", "심세은", "유수종"
    ]
}

s3_client = boto3.client(
    's3', 
    aws_access_key_id=AWS_ACCESS_KEY_ID, 
    aws_secret_access_key=AWS_SECRET_ACCESS_KEY, 
    region_name=AWS_REGION
)

@st.cache_data(ttl=60)
def get_all_s3_data(year_month_path):
    all_data = []
    prefix = f"data/{year_month_path}/"
    try:
        paginator = s3_client.get_paginator('list_objects_v2')
        for page in paginator.paginate(Bucket=S3_BUCKET_NAME, Prefix=prefix):
            if 'Contents' in page:
                for obj in page['Contents']:
                    if obj['Key'].endswith('.json'):
                        content = s3_client.get_object(Bucket=S3_BUCKET_NAME, Key=obj['Key'])['Body'].read().decode('utf-8')
                        all_data.extend(json.loads(content))
        
        df = pd.DataFrame(all_data)
        if not df.empty:
            if '비고' not in df.columns: df['비고'] = ""
            if '배달비_증빙URL' not in df.columns: df['배달비_증빙URL'] = ""
        return df
    except Exception as e:
        st.sidebar.error(f"S3 데이터 로드 오류: {e}")
        return pd.DataFrame()

def get_presigned_url(full_url):
    if not full_url or str(full_url).strip() in ["", "N/A", "nan"]: 
        return None
    try:
        pure_url = urllib.parse.unquote(full_url)
        key_start = pure_url.find("images/")
        if key_start != -1:
            s3_key = pure_url[key_start:]
            return s3_client.generate_presigned_url(
                'get_object',
                Params={'Bucket': S3_BUCKET_NAME, 'Key': s3_key},
                ExpiresIn=600
            )
    except: pass
    return None

# ==========================================
# 문서 생성 모듈 1: 개인별 엑셀 폼 
# ==========================================
def generate_excel_form(expense_items, user_name):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "경비지급신청서"

    border_thin = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))
    border_top_thin_bottom_double = Border(top=Side(style='thin'), bottom=Side(style='double'), left=Side(style='thin'), right=Side(style='thin'))
    
    align_center = Alignment(horizontal='center', vertical='center', wrap_text=True)
    align_left = Alignment(horizontal='left', vertical='center')
    align_right = Alignment(horizontal='right', vertical='center')
    font_bold = Font(bold=True)
    font_title = Font(name='맑은 고딕', size=16, bold=True)
    
    fill_c0c0c0 = PatternFill(start_color="C0C0C0", end_color="C0C0C0", fill_type="solid")

    def apply_border_to_range(range_string, border_style=border_thin):
        for row in ws[range_string]:
            for cell in row:
                cell.border = border_style

    ws.page_setup.paperSize = ws.PAPERSIZE_A4
    ws.page_setup.fitToPage = True
    ws.page_setup.fitToWidth = 1   
    ws.page_setup.fitToHeight = 0  
    ws.print_options.horizontalCentered = True
    ws.page_margins.left = 0.3
    ws.page_margins.right = 0.3

    ws.column_dimensions['A'].width = 12  
    ws.column_dimensions['B'].width = 22  
    ws.column_dimensions['C'].width = 16  
    ws.column_dimensions['D'].width = 13  
    ws.column_dimensions['E'].width = 7.0   
    for col in ['F', 'G', 'H', 'I']:
        ws.column_dimensions[col].width = 8 

    now = datetime.now()
    prev_month = 12 if now.month == 1 else now.month - 1
    target_month = f"{prev_month:02d}"

    ws.merge_cells('A1:D3')
    ws['A1'] = f"(주) 밀버스 {target_month}월 경비 지급신청"
    ws['A1'].font = font_title
    ws['A1'].alignment = align_left

    approvers = ["담당", "팀장", "본부장", "관리부"]
    
    apply_border_to_range('E1:I3', border_thin)

    ws.merge_cells('E1:E2')
    ws['E1'] = "결\n재" 
    ws['E1'].alignment = align_center
    
    ws['E3'] = "날 짜"
    ws['E3'].alignment = align_center
    
    ws.row_dimensions[1].height = 20
    ws.row_dimensions[2].height = 40
    ws.row_dimensions[3].height = 16 

    for idx, approver in enumerate(approvers):
        col_letter = chr(ord('F') + idx) 
        ws[f'{col_letter}1'] = approver
        ws[f'{col_letter}1'].alignment = align_center
        ws[f'{col_letter}2'] = "" 
        ws[f'{col_letter}3'] = "   /   " 
        ws[f'{col_letter}3'].alignment = align_center

    ws.merge_cells('A5:I5')
    ws['A5'] = f"사용자 : {user_name}"
    ws['A5'].font = font_bold
    ws['A5'].alignment = align_left

    total_amt = sum(item.get('_effective_cost', 0) for item in expense_items)
    
    ws.merge_cells('C7:D7')
    ws['C7'] = "청 구 액"
    ws['C7'].alignment = align_center
    ws['C7'].font = font_bold
    apply_border_to_range('C7:D7') 
    
    ws.merge_cells('E7:I7') 
    ws['E7'] = f"{total_amt:,} 원정"
    ws['E7'].alignment = align_right
    ws['E7'].font = font_bold
    apply_border_to_range('E7:I7') 

    headers = ["일 자", "사 용 처", "사 용 내 역", "금 액"]
    for col_num, header in enumerate(headers, 1):
        cell = ws.cell(row=9, column=col_num, value=header)
        cell.font = font_bold
        cell.alignment = align_center
        cell.fill = fill_c0c0c0
        cell.border = border_top_thin_bottom_double
        
    ws.merge_cells('E9:I9')
    ws['E9'] = "비 고 (slack 퇴근시간 등)"
    ws['E9'].font = font_bold
    ws['E9'].alignment = align_center
    ws['E9'].fill = fill_c0c0c0
    
    for c in range(5, 10):
        ws.cell(row=9, column=c).border = border_top_thin_bottom_double

    current_row = 10
    for item in expense_items:
        raw_date = item.get('결제일자', '')
        formatted_date = raw_date
        if raw_date:
            try:
                dt_obj = datetime.strptime(raw_date, "%Y-%m-%d")
                formatted_date = f"{dt_obj.month}/{dt_obj.day}"
            except Exception:
                pass
                
        ws.cell(row=current_row, column=1, value=formatted_date).alignment = align_center
        ws.cell(row=current_row, column=2, value=item.get('사용처', '')).alignment = align_left
        ws.cell(row=current_row, column=3, value=item.get('종류', '')).alignment = align_center
        
        amt_cell = ws.cell(row=current_row, column=4, value=item.get('_effective_cost', 0))
        amt_cell.alignment = align_right
        amt_cell.number_format = '#,##0'
        
        ws.merge_cells(start_row=current_row, start_column=5, end_row=current_row, end_column=9)
        ws.cell(row=current_row, column=5, value=item.get('비고', '')).alignment = align_left
        apply_border_to_range(f'A{current_row}:I{current_row}') 
        
        ws.row_dimensions[current_row].height = 21
        current_row += 1
        
        if item.get('배달비_이미지_display'):
            ws.cell(row=current_row, column=1, value=formatted_date).alignment = align_center
            delivery_shop_name = f"└ {item.get('사용처', '')} 배달비" 
            ws.cell(row=current_row, column=2, value=delivery_shop_name).alignment = align_left
            ws.cell(row=current_row, column=3, value=item.get('종류', '')).alignment = align_center
            
            del_amt_cell = ws.cell(row=current_row, column=4, value=0)
            del_amt_cell.alignment = align_right 
            del_amt_cell.number_format = '#,##0'
            
            ws.merge_cells(start_row=current_row, start_column=5, end_row=current_row, end_column=9)
            ws.cell(row=current_row, column=5, value="배달비 증빙 자료 첨부").alignment = align_left
            apply_border_to_range(f'A{current_row}:I{current_row}') 
            
            ws.row_dimensions[current_row].height = 21
            current_row += 1

    while current_row <= 27:
        ws.merge_cells(start_row=current_row, start_column=5, end_row=current_row, end_column=9)
        apply_border_to_range(f'A{current_row}:I{current_row}')
        ws.row_dimensions[current_row].height = 21
        current_row += 1

    ws.merge_cells(f'A{current_row}:C{current_row}')
    ws.cell(row=current_row, column=1, value="합        계").alignment = align_center
    ws.cell(row=current_row, column=1).font = font_bold
    
    tot_cell = ws.cell(row=current_row, column=4, value=total_amt)
    tot_cell.alignment = align_right
    tot_cell.font = font_bold
    tot_cell.number_format = '#,##0'
    
    ws.merge_cells(start_row=current_row, start_column=5, end_row=current_row, end_column=9)
    ws.cell(row=current_row, column=5, value="-").alignment = align_center
    
    for c in range(1, 10):
        ws.cell(row=current_row, column=c).border = border_top_thin_bottom_double

    current_row += 1 
    ws.merge_cells(f'A{current_row}:I{current_row}')
    ws.row_dimensions[current_row].height = 20

    current_row += 1 
    ws.merge_cells(f'A{current_row}:I{current_row}')
    ws.cell(row=current_row, column=1, value="상기 금액을 청구합니다.").alignment = align_center
    ws.row_dimensions[current_row].height = 15
    
    current_row += 1 
    today_str = datetime.now().strftime("%Y년 %m월 %d일")
    ws.merge_cells(f'A{current_row}:I{current_row}')
    ws.cell(row=current_row, column=1, value=today_str).alignment = align_center
    ws.row_dimensions[current_row].height = 15
    
    current_row += 1 
    ws.row_dimensions[current_row].height = 40
    ws.merge_cells(f'G{current_row}:I{current_row}') 

    current_dir = os.path.dirname(os.path.abspath(__file__))
    logo_path = os.path.join(current_dir, "logo.png")

    if os.path.exists(logo_path):
        try:
            with Image.open(logo_path) as pil_img:
                if pil_img.mode != 'RGBA':
                    pil_img = pil_img.convert('RGBA')
                img_byte_arr = io.BytesIO()
                pil_img.save(img_byte_arr, format='PNG')
                img_byte_arr.seek(0)
            
            logo_img = ExcelImage(img_byte_arr)
            logo_img.width = 160  
            logo_img.height = 40
            
            ws.add_image(logo_img, f"G{current_row}")
        except Exception as e:
            pass

    thick_border = Side(style='medium', color='000000')
    for r in range(1, current_row + 1):
        for c in range(1, 10):
            cell = ws.cell(row=r, column=c)
            b = cell.border
            if not b: b = Border()
            
            t = thick_border if r == 1 else b.top
            bot = thick_border if r == current_row else b.bottom
            l = thick_border if c == 1 else b.left
            ri = thick_border if c == 9 else b.right
            
            cell.border = Border(top=t, bottom=bot, left=l, right=ri, diagonal=b.diagonal, diagonalDown=b.diagonalDown)

    ws.sheet_view.showGridLines = False

    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return output

# ==========================================
# 문서 생성 모듈 2: 개인별 워드 영수증 모음 
# ==========================================
def generate_receipts_word(expense_items):
    receipt_imgs = []
    for item in expense_items:
        if item.get('image_display'):
            receipt_imgs.append(item['image_display'])
        if item.get('배달비_이미지_display'):
            receipt_imgs.append(item['배달비_이미지_display'])

    if not receipt_imgs:
        return None

    doc = docx.Document()
    
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(1.0)
        section.right_margin = Cm(1.0)
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(1.0)

    chunks = [receipt_imgs[i:i + 6] for i in range(0, len(receipt_imgs), 6)]

    for chunk_idx, chunk in enumerate(chunks):
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
        table.autofit = False

        for col in table.columns:
            for cell in col.cells:
                cell.width = Cm(9.5)

        for i in range(6):
            r_idx = i // 2 
            c_idx = i % 2  
            
            table.rows[r_idx].height = Cm(9.0)

            if i < len(chunk):
                img = chunk[i]
                img_stream = io.BytesIO()

                if img.mode != 'RGB':
                    img = img.convert('RGB')

                img.thumbnail((1200, 1200), Image.Resampling.LANCZOS)
                img.save(img_stream, format='JPEG', quality=85)
                img_stream.seek(0)

                cell = table.cell(r_idx, c_idx)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                img_w, img_h = img.size
                ratio = img_w / img_h
                
                target_w_cm, target_h_cm = 9.0, 8.6
                
                if ratio > (target_w_cm / target_h_cm): 
                    p.add_run().add_picture(img_stream, width=Cm(target_w_cm))
                else: 
                    p.add_run().add_picture(img_stream, height=Cm(target_h_cm))

        if chunk_idx < len(chunks) - 1:
            doc.add_page_break()

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# ==========================================
# 문서 생성 모듈 3: 월간 팀별 전체 집계표 엑셀
# ==========================================
def generate_team_aggregate_excel(df, team_name, year_month):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = f"월간 경비 집계 ({team_name})"

    align_c = Alignment(horizontal='center', vertical='center', wrap_text=True)
    align_r = Alignment(horizontal='right', vertical='center')
    font_bold = Font(name='맑은 고딕', bold=True)
    font_title = Font(name='맑은 고딕', size=18, bold=True)
    font_header = Font(name='맑은 고딕', size=10, bold=True)
    font_data = Font(name='맑은 고딕', size=10)
    
    fill_blue = PatternFill(start_color="BCE3F0", end_color="BCE3F0", fill_type="solid")
    fill_grey = PatternFill(start_color="E7E6E6", end_color="E7E6E6", fill_type="solid")
    fill_yellow = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")

    thin = Side(style='thin')
    medium = Side(style='medium')
    dotted = Side(style='dotted')

    ws.column_dimensions['A'].width = 5   
    ws.column_dimensions['B'].width = 10  
    ws.column_dimensions['C'].width = 14  
    for col in ['D','E','F','G','H','I','J','K','L','M']:
        ws.column_dimensions[col].width = 11.5
        
    ws.row_dimensions[1].height = 25
    ws.row_dimensions[2].height = 35
    ws.row_dimensions[3].height = 25
    ws.row_dimensions[4].height = 20
    ws.row_dimensions[5].height = 25
    ws.row_dimensions[6].height = 20

    ws.merge_cells('I1:I2')
    ws['I1'] = "결\n재"
    ws['I1'].alignment = align_c
    ws['I1'].font = font_header
    ws['I1'].border = Border(top=thin, bottom=thin, left=thin, right=thin)
    
    ws['I3'] = "날짜"
    ws['I3'].alignment = align_c
    ws['I3'].font = font_header
    ws['I3'].border = Border(top=thin, bottom=thin, left=thin, right=thin)
    
    approvers = ["팀장", "본부장", "관리본부", "대표이사"]
    for i, app in enumerate(approvers):
        col = 10 + i 
        ws.cell(row=1, column=col, value=app).alignment = align_c
        ws.cell(row=1, column=col).font = font_header
        ws.cell(row=1, column=col).border = Border(top=thin, bottom=thin, left=thin, right=thin)
        
        ws.cell(row=2, column=col, value="").border = Border(top=thin, bottom=thin, left=thin, right=thin)
        
        ws.cell(row=3, column=col, value="  /  ").alignment = align_c
        ws.cell(row=3, column=col).font = font_header
        ws.cell(row=3, column=col).border = Border(top=thin, bottom=thin, left=thin, right=thin)

    ws.merge_cells('A1:H3')
    y, m = year_month.split('/')
    ws['A1'] = f"㈜밀버스 {y}년 {int(m)}월 경비 사용내역"
    ws['A1'].font = font_title
    ws['A1'].alignment = align_c

    for r in range(4, 7):
        for c in range(1, 14):
            cell = ws.cell(row=r, column=c)
            cell.fill = fill_blue if c <= 3 else fill_grey
            cell.font = font_header
            cell.alignment = align_c
            cell.border = Border(top=thin, bottom=thin, left=thin, right=thin)

    ws.merge_cells('A4:A6')
    ws.merge_cells('B4:B6'); ws['B4'] = "이름"
    ws.merge_cells('C4:C6'); ws['C4'] = "경비\n사용금액"
    ws.merge_cells('D4:M4'); ws['D4'] = "상 세 내 역"

    cat_order = ["야근교통비", "야근식대", "외근교통비", "기타", "프로젝트비용"] 
    col_idx = 4
    for cat in ["야근교통비", "야근식대", "외근교통비", "기타", "프로젝트 비용"]: 
        ws.merge_cells(start_row=5, start_column=col_idx, end_row=5, end_column=col_idx+1)
        ws.cell(row=5, column=col_idx, value=cat)
        ws.cell(row=6, column=col_idx, value="개인카드")
        ws.cell(row=6, column=col_idx+1, value="법인카드")
        col_idx += 2

    ws['A4'].border = Border(top=thin, bottom=thin, left=thin, right=thin, diagonal=thin, diagonalDown=True)

    col_borders = {
        1: (thin, thin), 2: (thin, thin), 3: (thin, thin),
        4: (thin, dotted), 5: (dotted, thin), 6: (thin, dotted), 7: (dotted, thin),
        8: (thin, dotted), 9: (dotted, thin), 10: (thin, dotted), 11: (dotted, thin),
        12: (thin, dotted), 13: (dotted, thin)
    }

    pivot = df.pivot_table(index='이름', columns='항목', values='금액', aggfunc='sum', fill_value=0) if not df.empty else pd.DataFrame()
    for c in cat_order:
        if c not in pivot.columns: pivot[c] = 0

    if team_name in TEAM_MEMBERS:
        all_members = TEAM_MEMBERS[team_name]
        for m_name in all_members:
            if m_name not in pivot.index:
                pivot.loc[m_name] = 0 
        pivot = pivot.reindex(all_members).fillna(0)
    else:
        pivot = pivot.sort_index()

    current_row = 7
    total_sums = {c: 0 for c in cat_order}
    total_all = 0
    
    for idx, (name, row) in enumerate(pivot.iterrows(), 1):
        user_total = row.sum()
        total_all += user_total
        
        ws.cell(row=current_row, column=1, value=idx).alignment = align_c
        ws.cell(row=current_row, column=2, value=name).alignment = align_c
        
        c3 = ws.cell(row=current_row, column=3, value=user_total if user_total > 0 else "-")
        c3.alignment = align_r if user_total > 0 else align_c
        if user_total > 0: c3.number_format = '#,##0'
        
        col_idx = 4
        for cat in cat_order:
            val = row.get(cat, 0)
            total_sums[cat] += val
            
            p_cell = ws.cell(row=current_row, column=col_idx, value=val if val > 0 else "-")
            p_cell.alignment = align_r if val > 0 else align_c
            if val > 0: p_cell.number_format = '#,##0'
            
            b_cell = ws.cell(row=current_row, column=col_idx+1, value="-")
            b_cell.alignment = align_c
            
            col_idx += 2
            
        ws.row_dimensions[current_row].height = 20
        
        for c in range(1, 14):
            l_st, r_st = col_borders[c]
            ws.cell(row=current_row, column=c).border = Border(left=l_st, right=r_st, top=thin, bottom=thin)
            ws.cell(row=current_row, column=c).font = font_data
            
        current_row += 1

    ws.merge_cells(f'A{current_row}:B{current_row}')
    ws.cell(row=current_row, column=1, value="합   계").alignment = align_c
    
    ws.cell(row=current_row, column=3, value=total_all if total_all > 0 else "-").number_format = '#,##0'
    ws.cell(row=current_row, column=3).alignment = align_r if total_all > 0 else align_c
    
    col_idx = 4
    for cat in cat_order:
        tot_val = total_sums[cat]
        ws.cell(row=current_row, column=col_idx, value=tot_val if tot_val > 0 else "-").number_format = '#,##0'
        ws.cell(row=current_row, column=col_idx).alignment = align_r if tot_val > 0 else align_c
        ws.cell(row=current_row, column=col_idx+1, value="-").alignment = align_c
        col_idx += 2

    ws.row_dimensions[current_row].height = 22
    
    for c in range(1, 14):
        cell = ws.cell(row=current_row, column=c)
        cell.fill = fill_yellow
        cell.font = font_header
        l_st, r_st = col_borders[c]
        cell.border = Border(left=l_st, right=r_st, top=thin, bottom=thin)

    def apply_outer_thick_border(ws, min_row, max_row, min_col, max_col):
        for r in range(min_row, max_row + 1):
            for c in range(min_col, max_col + 1):
                b = ws.cell(row=r, column=c).border
                t = medium if r == min_row else b.top
                bot = medium if r == max_row else b.bottom
                l = medium if c == min_col else b.left
                ri = medium if c == max_col else b.right
                
                ws.cell(row=r, column=c).border = Border(top=t, bottom=bot, left=l, right=ri, diagonal=b.diagonal, diagonalDown=b.diagonalDown)

    apply_outer_thick_border(ws, 4, 6, 1, 13)
    apply_outer_thick_border(ws, 4, current_row, 1, 13)

    ws.sheet_view.showGridLines = False

    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return output

# ==========================================
# 2. 메인 화면 구성
# ==========================================
st.title("경비 정산 관리자 대시보드")
st.markdown("<p style='opacity: 0.7; font-size: 15px; margin-bottom: 2rem;'>임직원이 제출한 경비 내역과 증빙 자료를 검토하고 집계합니다.</p>", unsafe_allow_html=True)

with st.sidebar:
    st.markdown("<h3 style='margin-bottom: 0.5rem;'>조회 설정</h3>", unsafe_allow_html=True)
    target_date = st.date_input("조회 월 선택", value=datetime.today()) 
    year_month = target_date.strftime('%Y/%m')
    
    with st.spinner("데이터를 불러오는 중..."):
        raw_df = get_all_s3_data(year_month)

if not raw_df.empty:
    team_list = ["전체"] + sorted(raw_df['팀명'].dropna().unique().tolist())
    sel_team = st.sidebar.selectbox("팀 선택", team_list)
    display_df = raw_df if sel_team == "전체" else raw_df[raw_df['팀명'] == sel_team]

    # --- 1. 요약 집계 표 (Pivot) ---
    st.markdown(f"<h3 style='margin-bottom: 1rem;'>{year_month} [{sel_team}] 집계 현황</h3>", unsafe_allow_html=True)
    
    with st.container(border=True):
        pivot = display_df.pivot_table(index=['팀명', '이름'], columns='항목', values='금액', aggfunc='sum', fill_value=0) if not display_df.empty else pd.DataFrame()
        
        if sel_team == "전체":
            for t_name, members in TEAM_MEMBERS.items():
                for m in members:
                    if (t_name, m) not in pivot.index:
                        pivot.loc[(t_name, m), :] = 0
            pivot = pivot.sort_index()
        elif sel_team in TEAM_MEMBERS:
            members = TEAM_MEMBERS[sel_team]
            for m in members:
                if (sel_team, m) not in pivot.index:
                    pivot.loc[(sel_team, m), :] = 0
            pivot = pivot.reindex([(sel_team, m) for m in members]).fillna(0)

        for cat in FIXED_CATEGORIES:
            if cat not in pivot.columns: pivot[cat] = 0
        pivot = pivot[FIXED_CATEGORIES]
        pivot['합계'] = pivot.sum(axis=1)
        
        total_row = pivot.sum().to_frame().T
        total_row.index = pd.MultiIndex.from_tuples([('전체', '합계')])
        pivot = pd.concat([pivot, total_row])
        
        st.dataframe(pivot.style.format("{:,.0f}원"), use_container_width=True)
        
        st.write("")
        agg_excel_io = generate_team_aggregate_excel(display_df, sel_team, year_month)
        t_month = year_month.replace("/", "")
        st.download_button(
            label=f"[{sel_team}] 월간 집계표 엑셀 다운로드",
            data=agg_excel_io.getvalue(),
            file_name=f"{t_month}_월간경비집계_{sel_team}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

    st.markdown("<hr style='margin: 2rem 0; border-top: 1px solid rgba(148, 163, 184, 0.2);'>", unsafe_allow_html=True)

    # --- 2. 상세 내역 및 증빙 검토 ---
    st.markdown("<h3 style='margin-bottom: 1rem;'>상세 내역 및 증빙 검토</h3>", unsafe_allow_html=True)
    
    c1, c2 = st.columns([1, 3])
    submitted_users = sorted(display_df['이름'].dropna().unique())
    sel_user = c1.selectbox("조회 대상자 선택", submitted_users, label_visibility="collapsed")
    
    user_detail = display_df[display_df['이름'] == sel_user].sort_values(by='결제일자')
    
    user_proj_dates = user_detail['수행일자'].unique()
    proj_info = user_proj_dates[0] if len(user_proj_dates) > 0 else "정보 없음"
    c2.markdown(f"<div style='padding-top:8px; opacity:0.8; font-size:14px;'><b>프로젝트/수행 기간:</b> {proj_info}</div>", unsafe_allow_html=True)
    
    st.write("")
    
    h = st.columns([1.2, 1.2, 2.0, 2.5, 1.2, 1.2])
    headers = ["항목", "결제일자", "사용처", "비고 (동석자/기타)", "금액", "증빙 자료"]
    for i, name in enumerate(headers):
        h[i].markdown(f"<div style='font-size:13px; font-weight:600; opacity:0.7; padding-bottom:8px; border-bottom:2px solid rgba(148, 163, 184, 0.3); margin-bottom:8px;'>{name}</div>", unsafe_allow_html=True)

    for idx, row in user_detail.iterrows():
        with st.container(border=True):
            r = st.columns([1.2, 1.2, 2.0, 2.5, 1.2, 1.2])
            
            r[0].markdown(f"<div style='font-size:14px; font-weight:500; margin-top:6px;'>{row.get('항목', '-')}</div>", unsafe_allow_html=True)
            r[1].markdown(f"<div style='font-size:14px; opacity:0.8; margin-top:6px;'>{row.get('결제일자', '-')}</div>", unsafe_allow_html=True)
            r[2].markdown(f"<div style='font-size:14px; margin-top:6px;'>{row.get('사용처', '-')}</div>", unsafe_allow_html=True)
            
            note_text = row.get('비고', '')
            r[3].markdown(f"<div style='font-size:13px; opacity:0.6; margin-top:6px;'>{note_text if note_text else '-'}</div>", unsafe_allow_html=True)
            
            r[4].markdown(f"<div style='font-size:15px; font-weight:600; margin-top:6px;'>{row.get('금액', 0):,} 원</div>", unsafe_allow_html=True)
            
            with r[5]:
                btn_cols = st.columns(2)
                main_url = get_presigned_url(row.get('증빙URL'))
                if main_url:
                    with btn_cols[0]:
                        with st.popover("🧾"):
                            st.image(main_url, width=400)
                            
                del_url = get_presigned_url(row.get('배달비_증빙URL'))
                if del_url:
                    with btn_cols[1]:
                        with st.popover("🧾"):
                            st.image(del_url, width=400)

    # --- 3. 관리자 전용 개인 문서 일괄 다운로드 ---
    st.markdown("<hr style='margin: 2rem 0; border-top: 1px solid rgba(148, 163, 184, 0.2);'>", unsafe_allow_html=True)
    st.markdown("<h3 style='margin-bottom: 1rem;'>해당 직원 개인 정산서 다운로드</h3>", unsafe_allow_html=True)
    st.caption("클라우드(S3)에 저장된 증빙 이미지를 취합하여 사용자가 생성했던 것과 완벽히 동일한 서식의 엑셀과 워드 파일을 생성합니다.")

    doc_key = f"doc_{sel_user}_{year_month}"

    if st.button(f"'{sel_user}' 엑셀 및 증빙(Word) 생성하기", type="primary", use_container_width=True):
        with st.spinner("클라우드에서 증빙 이미지를 모아 문서를 생성하고 있습니다. (약 5~10초 소요)"):
            expense_items = []
            
            for _, row in user_detail.iterrows():
                main_img = None
                del_img = None

                m_url = get_presigned_url(row.get('증빙URL'))
                if m_url:
                    try:
                        res = requests.get(m_url, timeout=10)
                        if res.status_code == 200: main_img = Image.open(io.BytesIO(res.content))
                    except: pass

                d_url = get_presigned_url(row.get('배달비_증빙URL'))
                if d_url:
                    try:
                        res = requests.get(d_url, timeout=10)
                        if res.status_code == 200: del_img = Image.open(io.BytesIO(res.content))
                    except: pass

                expense_items.append({
                    '결제일자': row.get('결제일자', ''),
                    '사용처': row.get('사용처', ''),
                    '종류': row.get('항목', ''),
                    '_effective_cost': int(row.get('금액', 0)),
                    '비고': row.get('비고', ''),
                    'image_display': main_img,
                    '배달비_이미지_display': del_img
                })

            excel_io = generate_excel_form(expense_items, sel_user)
            word_io = generate_receipts_word(expense_items)

            st.session_state[f"excel_{doc_key}"] = excel_io.getvalue() if excel_io else None
            st.session_state[f"word_{doc_key}"] = word_io.getvalue() if word_io else None

    if f"excel_{doc_key}" in st.session_state:
        c_ex, c_wd = st.columns(2)
        target_m = year_month.replace("/", "")
        with c_ex:
            st.download_button("개인 엑셀 정산서 다운로드", data=st.session_state[f"excel_{doc_key}"], file_name=f"{sel_user}_경비지급신청서_{target_m}.xlsx", use_container_width=True)
        with c_wd:
            if st.session_state[f"word_{doc_key}"]:
                st.download_button("개인 증빙자료(Word) 다운로드", data=st.session_state[f"word_{doc_key}"], file_name=f"{sel_user}_증빙자료_{target_m}.docx", use_container_width=True)

else:
    st.info("해당 월에 제출된 정산 데이터가 없습니다.")