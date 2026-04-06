import streamlit as st
import pandas as pd
import boto3
import base64
import requests
import json
import io
import os
import time
import calendar
import uuid
import openpyxl
import zipfile # [추가] Word 내부 이미지 추출을 위한 모듈
from openpyxl.styles import Alignment, Font, Border, Side, PatternFill, Protection
from openpyxl.drawing.image import Image as ExcelImage
from datetime import datetime
from PIL import Image, ImageDraw

# Word 생성을 위한 모듈
import docx
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ==========================================
# 0. UI 설정 및 컴팩트 SaaS 디자인 CSS 적용
# ==========================================
st.set_page_config(page_title="경비 정산", layout="wide")

st.markdown("""
    <style>
    @import url("https://cdn.jsdelivr.net/gh/orioncactus/pretendard@v1.3.8/dist/web/variable/pretendardvariable.css");
    html, body, [class*="css"], .stMarkdown, .stText, button, input, select {
        font-family: 'Pretendard Variable', Pretendard, -apple-system, sans-serif !important;
    }
    
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    .stAppDeployButton {display:none;}
    header {background-color: transparent !important;}
    
    /* 우측 상단 툴바 완벽 숨김 */
    [data-testid="stHeaderActionElements"] {display: none !important;}
    [data-testid="stToolbar"] {visibility: hidden !important;}

    /* 카드 패딩 극한으로 축소 */
    div[data-testid="stVerticalBlockBorderWrapper"] {
        border-radius: 8px !important;
        box-shadow: rgba(0, 0, 0, 0.02) 0px 2px 4px !important;
        border: 1px solid rgba(148, 163, 184, 0.2) !important;
        background-color: rgba(255, 255, 255, 0.02) !important;
        padding: 4px 8px !important; 
        margin-bottom: 0px !important;
        transition: all 0.2s ease;
    }
    
    /* 기본 컬럼 갭 축소 */
    [data-testid="column"] > div {
        gap: 0.3rem !important;
    }
    
    /* 주요 버튼 */
    .stButton > button[kind="primary"] {
        background-color: #4f46e5;
        color: white;
        border: none;
        border-radius: 8px;
        font-weight: 600;
        letter-spacing: -0.3px;
        transition: all 0.2s ease;
    }
    .stButton > button[kind="primary"]:hover {
        background-color: #3730a3;
        box-shadow: 0 4px 12px rgba(79, 70, 229, 0.3);
    }
    
    /* 이모지 버튼 정중앙 정렬 및 여백 최소화 */
    .stButton > button[kind="secondary"], div[data-testid="stPopover"] > button {
        border-radius: 6px !important;
        font-weight: 600 !important;
        font-size: 16px !important; 
        border: 1px solid rgba(148, 163, 184, 0.3) !important;
        background-color: transparent !important;
        padding: 0px 6px !important;
        height: 36px !important;
        min-height: 36px !important; 
        display: flex !important;
        align-items: center !important;
        justify-content: center !important;
        margin: 0 !important;
    }
    .stButton > button[kind="secondary"]:hover, div[data-testid="stPopover"] > button:hover {
        background-color: rgba(148, 163, 184, 0.1) !important;
    }
    
    h1 { font-weight: 700 !important; letter-spacing: -1px; margin-bottom: 0px !important;}
    h3 { font-weight: 600 !important; letter-spacing: -0.5px; }
    
    /* 입력 폼 텍스트 잘림 방지 및 높이 통일 */
    div[data-baseweb="input"], div[data-baseweb="select"] {
        border-radius: 6px !important;
        border: none !important;
        background-color: rgba(148, 163, 184, 0.08) !important;
        box-shadow: inset 0 0 0 1px rgba(148, 163, 184, 0.2) !important;
        transition: all 0.2s ease;
        height: 36px !important;
        min-height: 36px !important; 
    }
    div[data-baseweb="input"]:focus-within, div[data-baseweb="select"]:focus-within {
        box-shadow: inset 0 0 0 2px #4f46e5 !important;
        background-color: rgba(148, 163, 184, 0.12) !important;
    }
    
    /* 일반 텍스트 입력창 여백 */
    div[data-baseweb="input"] > div > input {
        background-color: transparent !important;
        padding-left: 8px !important; 
        padding-right: 8px !important;
        padding-top: 2px !important;
        padding-bottom: 2px !important;
        font-size: 14px !important;
    }
    
    /* 드롭다운(Select) 내부 패딩을 0에 가깝게 줄여 글자 확보 */
    div[data-baseweb="select"] > div {
        background-color: transparent !important;
        padding-left: 4px !important; 
        padding-right: 0px !important;
        padding-top: 2px !important;
        padding-bottom: 2px !important;
        font-size: 14px !important;
    }
    
    /* Number Input의 + / - 버튼(스피너) 숨김 */
    [data-testid="stNumberInputStepUp"], 
    [data-testid="stNumberInputStepDown"] {
        display: none !important;
    }
    
    [data-testid="stFileUploadDropzone"] {
        border-radius: 6px !important;
        padding: 0.5rem !important;
        min-height: 60px !important;
    }
    
    div[role="radiogroup"] { gap: 0.5rem; }
    </style>
    """, unsafe_allow_html=True)

if 'expense_items' not in st.session_state: st.session_state.expense_items = []
if 'selected_cat' not in st.session_state: st.session_state.selected_cat = "야근식대"
if 'file_cat_map' not in st.session_state: st.session_state.file_cat_map = {}
if 'submitted' not in st.session_state: st.session_state.submitted = False
if 'uploader_key' not in st.session_state: st.session_state.uploader_key = 0

def change_category(cat_name):
    st.session_state.selected_cat = cat_name

# ==========================================
# 1. 유틸리티 함수
# ==========================================
def safe_int(value):
    try:
        if isinstance(value, str):
            clean_val = "".join(filter(lambda x: x.isdigit() or x == '-', value))
            return abs(int(clean_val)) if clean_val else 0
        return abs(int(value)) if value is not None else 0
    except: return 0

def analyze_receipt(uploaded_file, retries=1):
    try:
        api_key = st.secrets["OPENAI_API_KEY"]
    except KeyError:
        st.error("Secrets에 OPENAI_API_KEY가 없습니다.")
        return {"결제 날짜": "에러", "사용처": "키 없음", "합계 금액": 0}

    base64_image = base64.b64encode(uploaded_file.getvalue()).decode('utf-8')
    headers = {"Content-Type": "application/json", "Authorization": f"Bearer {api_key}"}
    
    prompt = """
    영수증 이미지에서 다음 3가지 정보를 반드시 추출하여 JSON 형식으로만 응답해.
    1. "결제 날짜": YYYY-MM-DD 형식. 
       * 주의: 한국 영수증은 'YY.MM.DD' 형식을 자주 사용해. (예: '26.03.19'는 2019년 3월 26일이 아니라 '2026년 3월 19일'이야. 무조건 맨 앞 두 자리를 연도(20YY)로 해석해!)
    2. "사용처": 상호명 추출. (택시/배달 앱의 경우 호출 옵션이나 가맹점 이름을 적어줘)
    3. "합계 금액": 최종 결제 금액 (숫자만).
    * 경고: 절대로 'None', 'null' 같은 문자열을 반환하지 마. 안 보이면 "미확인" 또는 0을 써.
    """
    
    payload = {
        "model": "gpt-4o-mini", 
        "temperature": 0.0, 
        "messages": [
            {"role": "system", "content": "너는 영수증 데이터를 기계처럼 정확하게 추출하는 시스템이야."},
            {"role": "user", "content": [{"type": "text", "text": prompt}, {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"}}]}
        ], 
        "response_format": { "type": "json_object" }
    }
    
    for attempt in range(retries + 1):
        try:
            response = requests.post("https://api.openai.com/v1/chat/completions", headers=headers, json=payload, timeout=45)
            res_data = json.loads(response.json()['choices'][0]['message']['content'])
            
            date_str = str(res_data.get("결제 날짜", "")).strip().lower()
            shop_str = str(res_data.get("사용처", "")).strip().lower()
            
            if (date_str in ["none", "null", ""] or shop_str in ["none", "null", ""]) and attempt < retries:
                time.sleep(1)
                continue
            
            if date_str in ["none", "null", "", "미확인"]: res_data["결제 날짜"] = datetime.now().strftime("%Y-%m-%d")
            else: res_data["결제 날짜"] = str(res_data.get("결제 날짜"))
            
            if shop_str in ["none", "null", "", "미확인"]: res_data["사용처"] = "미확인"
            else: res_data["사용처"] = str(res_data.get("사용처"))
            
            res_data["is_uncertain"] = (res_data.get("사용처") == "미확인" or safe_int(res_data.get("합계 금액")) == 0)
            return res_data
            
        except:
            if attempt < retries:
                time.sleep(1)
                continue
            break
            
    return {"결제 날짜": datetime.now().strftime("%Y-%m-%d"), "사용처": "분석 실패", "합계 금액": 0, "is_uncertain": True}

def save_to_s3(user_name, team_name, day_status, expense_items):
    now = datetime.now()
    date_path = now.strftime('%Y/%m')
    timestamp = now.strftime('%Y%m%d_%H%M%S')
    summary_list = []
    
    s3_bucket = st.secrets["S3_BUCKET_NAME"]
    aws_region = st.secrets["AWS_REGION"]
    s3_client = boto3.client('s3', aws_access_key_id=st.secrets["AWS_ACCESS_KEY_ID"], aws_secret_access_key=st.secrets["AWS_SECRET_ACCESS_KEY"], region_name=aws_region)

    for idx, item in enumerate(expense_items):
        final_amt = item.get('_effective_cost', 0)
        if final_amt == 0 and item['종류'] == "프로젝트비용":
            continue

        img_url = "N/A"
        del_img_url = "N/A"
        
        if item.get('image_display'):
            img_key = f"images/{date_path}/{team_name}/{user_name}_{timestamp}_{item.get('id', idx)}.png"
            img_byte_arr = io.BytesIO()
            item['image_display'].save(img_byte_arr, format='PNG')
            s3_client.put_object(Bucket=s3_bucket, Key=img_key, Body=img_byte_arr.getvalue(), ContentType='image/png')
            img_url = f"https://{s3_bucket}.s3.{aws_region}.amazonaws.com/{img_key}"
            
        if item.get('배달비_이미지_display'):
            del_img_key = f"images/{date_path}/{team_name}/{user_name}_{timestamp}_{item.get('id', idx)}_delivery.png"
            del_img_byte_arr = io.BytesIO()
            item['배달비_이미지_display'].save(del_img_byte_arr, format='PNG')
            s3_client.put_object(Bucket=s3_bucket, Key=del_img_key, Body=del_img_byte_arr.getvalue(), ContentType='image/png')
            del_img_url = f"https://{s3_bucket}.s3.{aws_region}.amazonaws.com/{del_img_key}"
            
        summary_list.append({
            "이름": user_name, "팀명": team_name, "항목": item['종류'], 
            "금액": final_amt, "결제일자": item['결제일자'], 
            "사용처": item['사용처'], "수행일자": day_status, 
            "비고": item.get('비고', ""), "증빙URL": img_url,
            "배달비_증빙URL": del_img_url if del_img_url != "N/A" else ""
        })
        
    s3_client.put_object(Bucket=s3_bucket, Key=f"data/{date_path}/{team_name}/{user_name}_{timestamp}.json", Body=json.dumps(summary_list, ensure_ascii=False).encode('utf-8'))
    return True

# ==========================================
# [엑셀] 폼 생성 함수
# ==========================================
def generate_excel_form(expense_items, user_name):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "경비지급신청서"

    border_thin = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))
    align_center = Alignment(horizontal='center', vertical='center', wrap_text=True)
    align_left = Alignment(horizontal='left', vertical='center')
    align_right = Alignment(horizontal='right', vertical='center')
    font_bold = Font(bold=True)
    font_title = Font(name='맑은 고딕', size=16, bold=True)

    def apply_border_to_range(range_string):
        for row in ws[range_string]:
            for cell in row:
                cell.border = border_thin

    ws.page_setup.paperSize = ws.PAPERSIZE_A4
    ws.page_setup.fitToPage = True
    ws.page_setup.fitToWidth = 1   
    ws.page_setup.fitToHeight = 0  
    ws.print_options.horizontalCentered = True
    ws.page_margins.left = 0.3
    ws.page_margins.right = 0.3

    ws.column_dimensions['A'].width = 12  
    ws.column_dimensions['B'].width = 22  
    ws.column_dimensions['C'].width = 16  
    ws.column_dimensions['D'].width = 13  
    ws.column_dimensions['E'].width = 4   
    for col in ['F', 'G', 'H', 'I']:
        ws.column_dimensions[col].width = 8 

    now = datetime.now()
    prev_month = 12 if now.month == 1 else now.month - 1
    target_month = f"{prev_month:02d}"

    ws.merge_cells('A1:D3')
    ws['A1'] = f"(주) 밀버스 {target_month}월 경비 지급신청"
    ws['A1'].font = font_title
    ws['A1'].alignment = align_left

    approvers = ["담당", "팀장", "본부장", "관리부"]
    ws.merge_cells('E1:E3')
    ws['E1'] = "결\n\n재"
    ws['E1'].alignment = align_center
    apply_border_to_range('E1:E3') 
    
    ws.row_dimensions[2].height = 45 

    for idx, approver in enumerate(approvers):
        col_letter = chr(ord('F') + idx) 
        ws[f'{col_letter}1'] = approver
        ws[f'{col_letter}1'].alignment = align_center
        ws[f'{col_letter}2'] = "" 
        ws[f'{col_letter}3'] = "   /   " 
        ws[f'{col_letter}3'].alignment = align_center
        apply_border_to_range(f'{col_letter}1:{col_letter}3') 

    ws.merge_cells('A5:I5')
    ws['A5'] = f"사용자 : {user_name}"
    ws['A5'].font = font_bold
    ws['A5'].alignment = align_left

    total_amt = sum(item.get('_effective_cost', 0) for item in expense_items)
    
    ws.merge_cells('C7:D7')
    ws['C7'] = "청 구 액"
    ws['C7'].alignment = align_center
    ws['C7'].font = font_bold
    apply_border_to_range('C7:D7') 
    
    ws.merge_cells('E7:I7') 
    ws['E7'] = f"{total_amt:,} 원정"
    ws['E7'].alignment = align_right
    ws['E7'].font = font_bold
    apply_border_to_range('E7:I7') 

    headers = ["일 자", "사 용 처", "사 용 내 역", "금 액"]
    for col_num, header in enumerate(headers, 1):
        cell = ws.cell(row=9, column=col_num, value=header)
        cell.font = font_bold
        cell.alignment = align_center
        cell.fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")
        cell.border = border_thin
        
    ws.merge_cells('E9:I9')
    ws['E9'] = "비 고 (slack 퇴근시간 등)"
    ws['E9'].font = font_bold
    ws['E9'].alignment = align_center
    ws['E9'].fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")
    apply_border_to_range('E9:I9')

    current_row = 10
    for item in expense_items:
        ws.cell(row=current_row, column=1, value=item.get('결제일자', '')).alignment = align_center
        ws.cell(row=current_row, column=2, value=item.get('사용처', '')).alignment = align_left
        ws.cell(row=current_row, column=3, value=item.get('종류', '')).alignment = align_center
        ws.cell(row=current_row, column=4, value=item.get('_effective_cost', 0)).alignment = align_right
        ws.merge_cells(start_row=current_row, start_column=5, end_row=current_row, end_column=9)
        ws.cell(row=current_row, column=5, value=item.get('비고', '')).alignment = align_left
        apply_border_to_range(f'A{current_row}:I{current_row}') 
        current_row += 1
        
        if item.get('배달비_이미지_display'):
            ws.cell(row=current_row, column=1, value=item.get('결제일자', '')).alignment = align_center
            delivery_shop_name = f"└ {item.get('사용처', '')} 배달비" 
            ws.cell(row=current_row, column=2, value=delivery_shop_name).alignment = align_left
            ws.cell(row=current_row, column=3, value=item.get('종류', '')).alignment = align_center
            ws.cell(row=current_row, column=4, value=0).alignment = align_right 
            ws.merge_cells(start_row=current_row, start_column=5, end_row=current_row, end_column=9)
            ws.cell(row=current_row, column=5, value="배달비 증빙 자료 첨부").alignment = align_left
            apply_border_to_range(f'A{current_row}:I{current_row}') 
            current_row += 1

    while current_row <= 22:
        ws.merge_cells(start_row=current_row, start_column=5, end_row=current_row, end_column=9)
        apply_border_to_range(f'A{current_row}:I{current_row}')
        current_row += 1

    ws.merge_cells(f'A{current_row}:C{current_row}')
    ws.cell(row=current_row, column=1, value="합        계").alignment = align_center
    ws.cell(row=current_row, column=1).font = font_bold
    ws.cell(row=current_row, column=4, value=total_amt).alignment = align_right
    ws.cell(row=current_row, column=4).font = font_bold
    
    ws.merge_cells(start_row=current_row, start_column=5, end_row=current_row, end_column=9)
    ws.cell(row=current_row, column=5, value="-").alignment = align_center
    apply_border_to_range(f'A{current_row}:I{current_row}')

    current_row += 2
    ws.merge_cells(f'A{current_row}:I{current_row}')
    ws.cell(row=current_row, column=1, value="상기 금액을 청구합니다.").alignment = align_center
    
    current_row += 2
    today_str = datetime.now().strftime("%Y년 %m월 %d일")
    ws.merge_cells(f'A{current_row}:I{current_row}')
    ws.cell(row=current_row, column=1, value=today_str).alignment = align_center

    ws.row_dimensions[current_row].height = 40 

    current_dir = os.path.dirname(os.path.abspath(__file__))
    logo_path = os.path.join(current_dir, "logo.png")

    if os.path.exists(logo_path):
        try:
            with Image.open(logo_path) as pil_img:
                if pil_img.mode != 'RGBA':
                    pil_img = pil_img.convert('RGBA')
                img_byte_arr = io.BytesIO()
                pil_img.save(img_byte_arr, format='PNG')
                img_byte_arr.seek(0)
            
            logo_img = ExcelImage(img_byte_arr)
            logo_img.width = 160  
            logo_img.height = 40
            
            ws.add_image(logo_img, f"G{current_row}")
        except Exception as e:
            print(f"로고 삽입 에러 발생: {e}")

    for row in ws['F2:I3']:
        for cell in row:
            cell.protection = Protection(locked=False)
            
    ws.protection.sheet = True

    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return output

# ==========================================
# 영수증 WORD 문서 생성 - 2x3 (총 6칸)
# ==========================================
def generate_receipts_word(expense_items):
    receipt_imgs = []
    for item in expense_items:
        if item.get('image_display'):
            receipt_imgs.append(item['image_display'])
        if item.get('배달비_이미지_display'):
            receipt_imgs.append(item['배달비_이미지_display'])

    if not receipt_imgs:
        return None

    doc = docx.Document()
    
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(1.0)
        section.right_margin = Cm(1.0)
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(1.0)

    chunks = [receipt_imgs[i:i + 6] for i in range(0, len(receipt_imgs), 6)]

    for chunk_idx, chunk in enumerate(chunks):
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
        table.autofit = False

        for col in table.columns:
            for cell in col.cells:
                cell.width = Cm(9.5)

        for i in range(6):
            r_idx = i // 2 
            c_idx = i % 2  
            
            table.rows[r_idx].height = Cm(9.0)

            if i < len(chunk):
                img = chunk[i]
                img_stream = io.BytesIO()

                if img.mode != 'RGB':
                    img = img.convert('RGB')

                img.thumbnail((1200, 1200), Image.Resampling.LANCZOS)
                img.save(img_stream, format='JPEG', quality=85)
                img_stream.seek(0)

                cell = table.cell(r_idx, c_idx)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                img_w, img_h = img.size
                ratio = img_w / img_h
                
                target_w_cm, target_h_cm = 9.0, 8.6
                
                if ratio > (target_w_cm / target_h_cm): 
                    p.add_run().add_picture(img_stream, width=Cm(target_w_cm))
                else: 
                    p.add_run().add_picture(img_stream, height=Cm(target_h_cm))

        if chunk_idx < len(chunks) - 1:
            doc.add_page_break()

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# [추가] 추출된 이미지 바이트 데이터를 분석 함수(analyze_receipt) 및 화면 표시에 호환되게 감싸는 클래스
class MockFile(io.BytesIO):
    def __init__(self, name, data):
        super().__init__(data)
        self.name = name
        self.type = "image/png" if name.lower().endswith('.png') else "image/jpeg"
    def getvalue(self):
        return super().getvalue()

# ==========================================
# 2. 메인 UI 및 사이드바 로직
# ==========================================
st.title("경비 정산")
st.markdown("<p style='color: #64748b; font-size: 15px; margin-bottom: 2rem;'>영수증 이미지 또는 <b>영수증이 포함된 Word(.docx) 파일</b>을 업로드하면 AI가 인식합니다.</p>", unsafe_allow_html=True)

with st.sidebar:
    st.markdown("<h3 style='margin-bottom: 0.5rem;'>제출자 정보</h3>", unsafe_allow_html=True)
    user_name = st.text_input("이름", placeholder="이름을 입력하세요")
    team_name = st.selectbox("소속 팀", ["관리본부", "DX1본부", "DX2본부", "CRM본부", "디자인팀", "ICT본부","영업본부","기타"])
    
    st.markdown("<hr style='margin: 1.5rem 0; border-top: 1px solid rgba(148, 163, 184, 0.2);'>", unsafe_allow_html=True)
    
    st.markdown("<h3 style='margin-bottom: 0.5rem;'>프로젝트 설정</h3>", unsafe_allow_html=True)
    project_type = st.radio("프로젝트 수행 여부", ["해당없음", "기간 선택"], horizontal=True, label_visibility="collapsed")
    
    max_project_cost = 0
    day_status = "해당없음"

    if project_type == "기간 선택":
        today = datetime.today()
        first_day = today.replace(day=1)
        dates = st.date_input("프로젝트 수행 기간", value=(first_day, today))
        
        if len(dates) == 2:
            start_date, end_date = dates
            working_days = (end_date - start_date).days + 1
            total_month_days = calendar.monthrange(start_date.year, start_date.month)[1]
            max_project_cost = int(200000 * (working_days / total_month_days))
            day_status = f"{start_date.strftime('%Y-%m-%d')} ~ {end_date.strftime('%Y-%m-%d')}"
            
            st.info(f"이번 달 배정 한도: {max_project_cost:,}원\n(근무 {working_days}일 / 해당월 {total_month_days}일)", icon="ℹ️")
        else:
            st.warning("달력에서 종료일을 선택해주세요.", icon="⚠️")

st.write("") 

categories = ["야근식대", "야근교통비", "외근교통비", "프로젝트비용", "기타"]
cols = st.columns(5)
for i, cat in enumerate(categories):
    cols[i].button(cat, use_container_width=True, type="primary" if st.session_state.selected_cat == cat else "secondary", on_click=change_category, args=(cat,))

st.write("") 

# [수정] 업로드 허용 파일 형식에 docx 추가
uploaded_files = st.file_uploader("증빙 자료(이미지 또는 Word 문서) 업로드", accept_multiple_files=True, type=['png', 'jpg', 'jpeg', 'docx'], key=f"receipt_uploader_{st.session_state.uploader_key}")

if uploaded_files:
    current_files = [f.name for f in uploaded_files]
    keys_to_remove = [k for k in st.session_state.file_cat_map.keys() if k not in current_files]
    for k in keys_to_remove:
        st.session_state.file_cat_map.pop(k)
        
    for f in uploaded_files:
        if f.name not in st.session_state.file_cat_map:
            st.session_state.file_cat_map[f.name] = st.session_state.selected_cat
            
    st.caption("분류 대기 항목: " + " | ".join([f"{f.name} ({st.session_state.file_cat_map[f.name]})" for f in uploaded_files]))
else:
    st.session_state.file_cat_map.clear()

if uploaded_files and st.button("파일 자동 입력 시작", type="primary", use_container_width=True):
    st.session_state.submitted = False 
    
    # ==========================================
    # [핵심] 일반 이미지와 Word 문서 속 이미지를 평탄화(Flatten) 처리
    # ==========================================
    processed_images = []
    
    for f in uploaded_files:
        if f.name.lower().endswith('.docx'):
            try:
                # ZipFile을 이용해 Word 내부 미디어 폴더 강제 추출
                with zipfile.ZipFile(f) as docx_zip:
                    media_files = [name for name in docx_zip.namelist() if name.startswith('word/media/') and name.lower().endswith(('.png', '.jpg', '.jpeg'))]
                    for img_name in media_files:
                        img_bytes = docx_zip.read(img_name)
                        
                        # Word 문서 내부의 원본 이미지 이름(image1.png 등)을 추출
                        extracted_name = f"{f.name}_{img_name.split('/')[-1]}"
                        
                        # 추출된 바이트 데이터를 처리 가능한 파일 객체 형태로 포장
                        mock_file = MockFile(extracted_name, img_bytes)
                        processed_images.append(mock_file)
            except Exception as e:
                st.error(f"{f.name} 파일에서 이미지를 추출하는 데 실패했습니다.")
        else:
            # 일반 이미지 파일은 그대로 리스트에 추가
            processed_images.append(f)
            
    total_files = len(processed_images)
    
    if total_files == 0:
        st.warning("분석할 수 있는 이미지가 없습니다.")
        st.stop()
        
    progress_bar = st.progress(0, text=f"총 {total_files}건의 이미지를 분석 중입니다...")
    
    for i, img_obj in enumerate(processed_images):
        # 원본 파일명 기반 매핑 (Word 추출 이미지는 현재 선택된 카테고리를 따름)
        original_file_name = img_obj.name.split('_word/media/')[0] if 'word/media/' in img_obj.name else img_obj.name
        assigned_cat = st.session_state.file_cat_map.get(original_file_name, st.session_state.selected_cat)
        
        # 분석 함수 호출 (img_obj 내부 데이터를 넘김)
        res = analyze_receipt(img_obj) 
        
        img = Image.open(img_obj)
        img.thumbnail((1500, 1500), Image.Resampling.LANCZOS)
        
        st.session_state.expense_items.append({
            "id": str(uuid.uuid4()), 
            "종류": assigned_cat, 
            "결제일자": res.get("결제 날짜"), 
            "사용처": res.get("사용처"), 
            "인식금액": safe_int(res.get("합계 금액")), 
            "비고": "", 
            "image_display": img, 
            "배달비_이미지_display": None, 
            "is_uncertain": res.get("is_uncertain", False)
        })
        
        if i < total_files - 1:
            time.sleep(3) # AI 과부하 방지 3초 딜레이
            
        progress_percentage = int(((i + 1) / total_files) * 100)
        progress_bar.progress(progress_percentage, text=f"총 {total_files}건 중 {i+1}건 완료...")
        
    st.session_state.expense_items.sort(key=lambda x: str(x.get('결제일자', '')))
    st.session_state.file_cat_map = {} 
    st.session_state.uploader_key += 1 
    time.sleep(0.5)
    progress_bar.empty()
    st.rerun()

# ==========================================
# 3. 리스트 표시, 자동 절사 및 제출 로직
# ==========================================
if st.session_state.expense_items:
    st.markdown("<hr style='margin: 2rem 0; border-top: 1px solid rgba(148, 163, 184, 0.2);'>", unsafe_allow_html=True)
    limit = max_project_cost if project_type == "기간 선택" else 0
    current_proj_total = 0
    
    limit_exceeded = False
    for i in st.session_state.expense_items:
        if i['종류'] == "프로젝트비용":
            if current_proj_total + i['인식금액'] > limit:
                limit_exceeded = True
                break
            current_proj_total += i['인식금액']

    if limit_exceeded and not st.session_state.submitted:
        if limit == 0:
            st.error("프로젝트 기간이 설정되지 않아 관련 비용이 제외 처리되었습니다.", icon="🚨")
        else:
            st.warning(f"프로젝트 한도({limit:,}원) 초과분이 자동으로 절사되었습니다.", icon="⚠️")

    st.markdown("<h3 style='margin-bottom: 1rem;'>정산 내역 확인</h3>", unsafe_allow_html=True)
    
    if st.session_state.submitted:
        st.success("해당 내역이 시스템에 성공적으로 등록되었습니다.", icon="✅")

    current_proj_total = 0 
    
    for idx, item in enumerate(st.session_state.expense_items):
        if 'id' not in item: item['id'] = str(uuid.uuid4())
        uid = item['id']

        with st.container(border=True):
            input_cost = item['인식금액']
            is_high_cost_meal = (item['종류'] == "야근식대" and input_cost >= 15000)

            r1 = st.columns([1.7, 1.1, 1.6, 1.2, 1.0, 1.5, 0.4, 0.4], vertical_alignment="center")
            
            item['종류'] = r1[0].selectbox(f"cat_{uid}", categories, index=categories.index(item['종류']), label_visibility="collapsed", disabled=st.session_state.submitted)
            item['결제일자'] = r1[1].text_input(f"dt_{uid}", item['결제일자'], label_visibility="collapsed", disabled=st.session_state.submitted)
            item['사용처'] = r1[2].text_input(f"vn_{uid}", item['사용처'], label_visibility="collapsed", disabled=st.session_state.submitted)
            item['인식금액'] = r1[3].number_input(f"am_{uid}", value=safe_int(item['인식금액']), step=100, label_visibility="collapsed", disabled=st.session_state.submitted)
            
            effective_cost = input_cost
            base_html = "<div style='display:flex; flex-direction:column; justify-content:center; height:36px; line-height:1.2;'>"
            
            if item['종류'] == "프로젝트비용":
                if limit == 0:
                    effective_cost = 0
                    status_html = f"{base_html}<del style='color:#94a3b8; font-size:12px;'>{input_cost:,}</del><span style='color:#ef4444; font-size:11px; font-weight:600;'>미설정 제외</span></div>"
                elif current_proj_total >= limit:
                    effective_cost = 0
                    status_html = f"{base_html}<del style='color:#94a3b8; font-size:12px;'>{input_cost:,}</del><span style='color:#ef4444; font-size:11px; font-weight:600;'>한도 초과</span></div>"
                elif current_proj_total + input_cost > limit:
                    effective_cost = limit - current_proj_total
                    current_proj_total = limit
                    cut_amt = input_cost - effective_cost
                    status_html = f"{base_html}<span style='font-size:14px; font-weight:600;'>{effective_cost:,} 원</span><span style='color:#ef4444; font-size:10px; font-weight:600;'>-{cut_amt:,}원 절사</span></div>"
                else:
                    effective_cost = input_cost
                    current_proj_total += effective_cost
                    status_html = f"{base_html}<span style='font-size:14px; font-weight:600;'>{effective_cost:,} 원</span></div>"
            else:
                status_html = f"{base_html}<span style='font-size:14px; font-weight:600;'>{effective_cost:,} 원</span></div>"
                    
            item['_effective_cost'] = effective_cost
            r1[4].markdown(status_html, unsafe_allow_html=True)
            
            if is_high_cost_meal:
                r1[5].markdown(f"{base_html}<span style='color:#ef4444; font-size:12px; font-weight:600;'>하단 증빙 필요 ↓</span></div>", unsafe_allow_html=True)
            else:
                item['배달비_이미지_display'] = None
                item['비고'] = r1[5].text_input("자유비고", value=item.get('비고', ''), placeholder="비고(선택)", key=f"note_free_{uid}", label_visibility="collapsed", disabled=st.session_state.submitted)

            with r1[6]:
                with st.popover("🧾"): 
                    st.image(item['image_display'], width=400)
            if r1[7].button("🗑️", key=f"del_{uid}", disabled=st.session_state.submitted):
                st.session_state.expense_items = [x for x in st.session_state.expense_items if x['id'] != uid]
                st.rerun()

            if is_high_cost_meal:
                st.markdown("<hr style='margin: 0.2rem 0 0.4rem 0; border-top: 1px solid rgba(79, 70, 229, 0.2);'>", unsafe_allow_html=True)
                
                c_sub1, c_sub2 = st.columns([1, 1], vertical_alignment="center")
                
                with c_sub1:
                    item['비고'] = st.text_input("동석자 및 비고", value=item.get('비고', ''), placeholder="동석자 정보 (예: 홍길동, 김철수) 또는 비고 입력", key=f"note_{uid}", label_visibility="collapsed", disabled=st.session_state.submitted)
                
                with c_sub2:
                    d1, d2 = st.columns([4, 1], vertical_alignment="center")
                    
                    del_file = d1.file_uploader("배달비 영수증 첨부", type=["png", "jpg", "jpeg"], key=f"del_file_{uid}", label_visibility="collapsed", disabled=st.session_state.submitted)
                    
                    if del_file:
                        del_img = Image.open(del_file)
                        del_img.thumbnail((1500, 1500), Image.Resampling.LANCZOS)
                        item['배달비_이미지_display'] = del_img
                    else:
                        item['배달비_이미지_display'] = None
                        
                    with d2:
                        if item.get('배달비_이미지_display'):
                            with st.popover("🧾"): 
                                st.image(item['배달비_이미지_display'], width=400)

    st.write("")
    
    valid_items = [
        item for item in st.session_state.expense_items 
        if not (item['종류'] == "프로젝트비용" and item.get('_effective_cost', 0) == 0)
    ]
    final_sorted_items = sorted(valid_items, key=lambda x: str(x.get('결제일자', '')))
    
    col_submit, col_excel, col_word = st.columns([1.2, 1, 1])
    
    now = datetime.now()
    if now.month == 1:
        target_m = f"{now.year - 1}12"
    else:
        target_m = f"{now.year}{now.month - 1:02d}"
    
    with col_submit:
        if not st.session_state.submitted:
            if st.button("최종 제출하기", type="primary", use_container_width=True):
                if not user_name: st.error("제출자 이름을 확인해주세요.", icon="🚨")
                elif project_type == "기간 선택" and max_project_cost == 0:
                    st.error("달력에서 프로젝트 종료일을 확인해주세요.", icon="🚨")
                else:
                    with st.spinner("서버에 데이터를 등록하고 있습니다..."):
                        if save_to_s3(user_name, team_name, day_status, final_sorted_items):
                            st.toast('정산 내역이 성공적으로 등록되었습니다.', icon='✔️')
                            st.session_state.submitted = True 
                            time.sleep(1) 
                            st.rerun()
        else:
            if st.button("새 정산 작성하기", use_container_width=True):
                st.session_state.expense_items = []
                st.session_state.submitted = False
                st.session_state.uploader_key += 1
                st.rerun()
                
    with col_excel:
        if final_sorted_items:
            excel_file = generate_excel_form(final_sorted_items, user_name)
            
            st.download_button(
                label="엑셀 양식 다운로드",
                data=excel_file,
                file_name=f"{user_name}_경비지급신청서_{target_m}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True
            )
            
    with col_word:
        if final_sorted_items:
            word_file = generate_receipts_word(final_sorted_items)
            
            if word_file:
                st.download_button(
                    label="영수증 모음 다운로드",
                    data=word_file,
                    file_name=f"{user_name}_증빙자료_{target_m}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )